from pathlib import Path
import os
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = Path(os.environ.get("REPORT_DOCS_DIR", ROOT / "docs"))
DOCX = Path(os.environ.get("REPORT_DOCX_PATH", DOCS_DIR / "CS599_大作业报告_程辉高.docx"))


REQUIRED_COVER_FIELDS = [
    "课程名称",
    "项目名称",
    "方向",
    "学号",
    "姓名",
    "专业",
    "指导教师",
    "提交日期",
]

REQUIRED_HEADINGS = [
    "一、选题背景与设计思想",
    "二、Specs 规格文档",
    "三、系统架构与设计",
    "四、关键实现与代码展示",
    "五、测试与评估",
    "六、系统升级与扩展",
    "七、课程总结",
    "附录 A：评分标准对照",
    "附录 B：时间节点与提交说明",
]

FORBIDDEN_PLACEHOLDERS = [
    "待填写",
    "请替换",
    "2025XXXXXXXX",
]


def main():
    document = Document(DOCX)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    all_text = text + "\n" + table_text

    missing_fields = [field for field in REQUIRED_COVER_FIELDS if field not in all_text]
    missing_headings = [heading for heading in REQUIRED_HEADINGS if heading not in text]
    leftover_placeholders = [token for token in FORBIDDEN_PLACEHOLDERS if token in all_text]
    heading_count = sum(1 for paragraph in document.paragraphs if paragraph.style.name.startswith("Heading"))

    with ZipFile(DOCX) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")

    has_toc_field = 'TOC \\o "1-3"' in document_xml
    has_hyperlink_flag = "\\h" in document_xml

    result = {
        "docx": str(DOCX),
        "tables": len(document.tables),
        "paragraphs": len(document.paragraphs),
        "heading_count": heading_count,
        "has_toc_field": has_toc_field,
        "has_toc_hyperlinks": has_hyperlink_flag,
        "missing_cover_fields": missing_fields,
        "missing_required_headings": missing_headings,
        "leftover_placeholders": leftover_placeholders,
        "passed": not missing_fields
        and not missing_headings
        and not leftover_placeholders
        and heading_count >= 8
        and has_toc_field,
    }
    print(result)
    if not result["passed"]:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
