import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = Path(os.environ.get("REPORT_DOCS_DIR", ROOT / "docs"))
OUT = Path(os.environ.get("REPORT_DOCX_OUT", DOCS_DIR / "CS599_大作业报告_程辉高.docx"))


STUDENT_ID = os.environ.get("STUDENT_ID", "2025302979")
STUDENT_NAME = os.environ.get("STUDENT_NAME", "程辉高")
STUDENT_MAJOR = os.environ.get("STUDENT_MAJOR", "计算机技术 / 软件工程")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("CS599 企业级应用软件设计与开发期末大作业  |  第 ")
    set_run_font(run, size=9, color=RGBColor(102, 112, 133))
    field_run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_char_begin)
    field_run._r.append(instr_text)
    field_run._r.append(fld_char_end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=RGBColor(102, 112, 133))


def add_toc(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    hint = document.add_paragraph("提示：首次打开 Word 文档后，右键目录并选择“更新域”即可刷新页码；左侧导航窗格可按标题跳转。")
    hint.style = document.styles["Body Text"]


def set_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 77, 120)),
        ("Heading 2", 13, RGBColor(37, 99, 169)),
        ("Heading 3", 12, RGBColor(23, 32, 42)),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    for name in ("Title", "Subtitle", "Body Text"):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title(document, text, subtitle=None):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(23, 32, 42)
    if subtitle:
        s = document.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = s.add_run(subtitle)
        sr.font.name = "Microsoft YaHei"
        sr._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        sr.font.size = Pt(13)
        sr.font.color.rgb = RGBColor(71, 84, 103)


def add_cover(document):
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    document.add_paragraph()
    add_title(document, "《企业级应用软件设计与开发》期末大作业报告", "企业项目协同系统的智能运营 Agent 改造")
    document.add_paragraph()

    cover_rows = [
        ("课程名称", "企业级应用软件设计与开发"),
        ("项目名称", "企业项目协同系统的智能运营 Agent 改造"),
        ("方向", "方向二：企业级应用软件的 Agent 改造"),
        ("学号", STUDENT_ID),
        ("姓名", STUDENT_NAME),
        ("专业", STUDENT_MAJOR),
        ("指导教师", "戚欣"),
        ("提交日期", "2026 年 6 月 22 日"),
    ]
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(11)
    set_cell_text(table.rows[0].cells[0], "字段", True)
    set_cell_text(table.rows[0].cells[1], "内容", True)
    set_cell_shading(table.rows[0].cells[0], "E8EEF5")
    set_cell_shading(table.rows[0].cells[1], "E8EEF5")
    for key, value in cover_rows:
        row = table.add_row()
        set_cell_text(row.cells[0], key, True)
        set_cell_text(row.cells[1], value)

    document.add_paragraph()
    note = document.add_paragraph("报告版本：最终提交版｜交付形式：GitHub 仓库 + Word/PDF 报告")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(102, 112, 133)
    document.add_page_break()


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            set_cell_text(cells[i], str(value))
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    document.add_paragraph()
    return table


def add_code_block(document, code):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F7F9FC")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for line in code.strip("\n").splitlines():
        run = paragraph.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)
    document.add_paragraph()


def add_diagram_table(document, title, headers, rows):
    cap = document.add_paragraph(title)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].bold = True
    cap.runs[0].font.size = Pt(10)
    add_table(document, headers, rows)


def add_screenshot_placeholder(document, title, description):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "FFF7E6")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title + "\n")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    detail = paragraph.add_run(description)
    detail.font.name = "Microsoft YaHei"
    detail._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    detail.font.size = Pt(9)
    detail.font.color.rgb = RGBColor(122, 90, 0)
    document.add_paragraph()


def paragraph(document, text):
    p = document.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(document, text):
    p = document.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)


def numbered(document, text):
    p = document.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(3)


def build_document():
    document = Document()
    document.core_properties.title = "CS599 大作业报告：企业项目协同系统的智能运营 Agent 改造"
    document.core_properties.author = STUDENT_NAME
    document.core_properties.subject = "企业级应用软件设计与开发"
    document.core_properties.keywords = "CS599, Agentic AI, SDD, MCP, Agentic RAG"
    set_styles(document)
    add_cover(document)

    document.add_heading("目录", level=1)
    add_toc(document)
    document.add_page_break()

    document.add_heading("一、选题背景与设计思想", level=1)
    paragraph(
        document,
        "本项目选择方向二：企业级应用软件的 Agent 改造。项目先构建一个模拟企业项目协同系统，覆盖客户、项目、合同、工单、审批、用户权限和审计日志等典型企业后台模块，再在其上叠加智能运营 Agent，使传统后台从“用户查表和填表”升级为“用户提出目标，Agent 编排工具完成任务”。",
    )
    paragraph(
        document,
        "原始系统的主要问题不是没有数据，而是缺少把数据转化为行动的智能层：项目风险需要人工跨表判断，工单派单依赖项目经理经验，合同审批需要人工切换页面处理，经营周报需要复制多个模块指标。这些问题在真实企业软件中普遍存在，适合作为 Agentic AI 改造场景。",
    )
    add_table(
        document,
        ["分析维度", "改造前", "改造后"],
        [
            ["交互方式", "菜单、表格、人工筛选", "自然语言任务驱动"],
            ["风险识别", "人工跨客户、项目、合同、工单判断", "Agent 汇聚多表数据并解释风险原因"],
            ["工单处理", "项目经理人工查看人员负载", "Tool Use 自动选择处理人并更新状态"],
            ["合同审批", "人工进入审批页面处理", "Agent 调用审批工具并写入审计"],
            ["经营报告", "手工统计指标", "一键生成经营周报"],
        ],
        [3.2, 6.2, 6.2],
    )
    paragraph(
        document,
        "技术路线采用 SDD 规格驱动开发：先定义 Product Spec，明确用户角色、痛点和验收标准；再定义 Architecture Spec，设计 Web API、Agent、工具层、记忆层和数据层；最后定义 API Spec，固化接口、权限和工具调用约束。实现层面采用 Python + SQLite 构建可离线演示系统，并优先支持免费 DeepSeek 兼容 API Tool Calling Planner。",
    )

    document.add_heading("二、Specs 规格文档", level=1)
    paragraph(document, "规格文档位于 docs/specs/，覆盖 Product Spec、Architecture Spec 与 API Spec，构成 SDD 核心交付。")
    add_table(
        document,
        ["规格文档", "核心内容", "可执行性体现"],
        [
            ["Product Spec", "项目定位、用户角色、痛点、用户故事、验收标准", "每个验收项映射到业务模块、Agent 工具或测试脚本"],
            ["Architecture Spec", "总体架构、Agent 交互流程、数据流、改造前后对比", "架构层级对应 app/、app/agent/、scripts/、docs/"],
            ["API Spec", "鉴权 API、业务资源 API、Agent API、MCP/Tool API、权限矩阵", "接口已在 app/server.py 中实现并由 tests.py 验证"],
        ],
        [3.1, 6.2, 6.5],
    )
    paragraph(document, "Product Spec 规定了六个核心验收项：用户登录与业务模块访问、经营周报生成、项目风险分析、Agent 工具写操作、运行轨迹留痕、测试和 benchmark 可运行。实现时每个验收项均可通过代码、界面或脚本验证。")

    document.add_heading("三、系统架构与设计", level=1)
    paragraph(document, "系统采用分层架构，将传统企业后台与 Agentic AI 层清晰分离。业务模块负责稳定的数据管理，Agent 层负责自然语言任务规划、工具调用、记忆和可观测性。")
    add_diagram_table(
        document,
        "图 3-1 系统总体架构图",
        ["层级", "模块", "职责"],
        [
            ["表现层", "Web 前端 app/static", "登录、看板、业务表格、智能助手、工具轨迹展示"],
            ["控制层", "Python HTTP API app/server.py", "路由、Session、RBAC、静态资源服务"],
            ["业务层", "客户/项目/合同/工单/审批", "提供企业系统基础数据与流程"],
            ["Agent 层", "EnterpriseAgent", "规划任务、调用工具、组织最终答复"],
            ["工具层", "Tool Registry", "Function Calling 工具定义与执行"],
            ["数据层", "SQLite", "业务数据、知识库、记忆、run/step 轨迹、审计日志"],
        ],
    )
    add_diagram_table(
        document,
        "图 3-2 Agent 交互流程",
        ["步骤", "输入/动作", "输出"],
        [
            ["1", "用户输入自然语言任务", "message"],
            ["2", "Agent 创建 run 并读取最近记忆", "run_id + memory"],
            ["3", "Planner 生成工具调用计划", "tool_name + tool_args"],
            ["4", "Tool Registry 调用业务工具", "observation"],
            ["5", "保存 agent_steps 与 audit_logs", "可观测轨迹"],
            ["6", "Answer Composer 生成答复", "最终自然语言答案"],
        ],
    )
    add_diagram_table(
        document,
        "图 3-3 数据流设计",
        ["数据类型", "来源", "去向"],
        [
            ["业务数据", "customers/projects/contracts/tickets", "看板、风险分析、周报生成"],
            ["规则知识", "knowledge_base", "search_knowledge 工具与 Agentic RAG"],
            ["对话记忆", "agent_memories", "Planner 上下文与连续任务"],
            ["运行轨迹", "agent_runs/agent_steps", "前端工具调用轨迹和评估分析"],
            ["审计数据", "audit_logs", "安全合规和可解释性展示"],
        ],
    )

    document.add_heading("四、关键实现与代码展示", level=1)
    paragraph(document, "Agent 核心循环位于 app/agent/agent.py。每次运行都会读取记忆、创建 run、生成工具计划、逐步调用工具、记录 observation、写入 assistant 记忆并返回最终答案。")
    add_code_block(
        document,
        """
def run_stream(self, message):
    memory = load_recent_memory(self.connection, self.user["id"])
    run_id = create_run(self.connection, self.user["id"], message)
    remember(self.connection, self.user["id"], "user", message)
    plan = self.llm_planner.plan(message, memory) or self.fallback_planner.plan(message, memory)
    for index, (thought, tool_name, tool_args) in enumerate(plan, start=1):
        observation = call_tool(tool_name, tool_args, self.connection, self.user)
        record_step(self.connection, run_id, index, thought, tool_name, tool_args, observation)
    answer = self.compose_answer(message, observations)
    finish_run(self.connection, run_id, answer)
""",
    )
    paragraph(document, "工具定义位于 app/agent/tools.py，核心工具覆盖经营指标读取、项目查询、工单查询、工单自动派单、风险分析、周报生成、合同审批和知识库检索。")
    add_table(
        document,
        ["工具", "能力", "对应课程技术点"],
        [
            ["get_dashboard_metrics", "读取经营看板指标", "Tool Use"],
            ["analyze_project_risks", "融合项目、合同、工单并计算风险", "多步骤推理"],
            ["route_ticket", "按负载自动派单并更新业务数据", "Function Calling 写操作"],
            ["approve_contract", "审批或驳回合同", "业务流程自动化"],
            ["search_knowledge", "检索企业规则知识", "Agentic RAG"],
            ["generate_weekly_report", "生成经营周报", "Agent 综合输出"],
        ],
        [4.0, 6.5, 5.0],
    )
    paragraph(document, "系统还提供 MCP stdio JSON-RPC 工具服务 scripts/mcp_server.py，使外部 Agent 客户端能够发现并调用本系统工具；同时提供 app/agent/llm.py，可在配置 FREE_DEEPSEEK_API_KEY 或 SILICONFLOW_API_KEY 后调用免费 DeepSeek 兼容 API 生成真实工具调用计划，官方 DEEPSEEK_API_KEY 仅作为备用。")
    add_screenshot_placeholder(document, "图 4-1 AI IDE 使用截图位置", "提交最终版前建议替换为 Trae CN / Codex / VS Code 中编写 Specs、实现 Agent 工具、运行测试的截图。")

    document.add_heading("五、测试与评估", level=1)
    paragraph(document, "项目提供功能测试、Agent 行为评估和 Web smoke test，避免只凭界面主观判断。")
    add_table(
        document,
        ["验证方式", "命令", "验证内容"],
        [
            ["功能测试", "python tests.py", "密码哈希、Agent 风险分析、登录 API、看板 API、周报 Agent API"],
            ["Benchmark", "python scripts/benchmark_agent.py", "工具调用命中率、关键词命中率、通过状态"],
            ["Web Smoke", "python scripts/smoke_web.py", "临时启动服务，验证首页、登录、看板和 Agent API"],
            ["Demo API", "python scripts/demo_api.py", "对已启动服务执行命令行演示"],
        ],
        [3.0, 5.3, 7.2],
    )
    add_table(
        document,
        ["Benchmark 用例", "Prompt", "期望工具"],
        [
            ["risk-analysis", "分析当前高风险项目并给出预警", "search_knowledge, analyze_project_risks"],
            ["weekly-report", "生成本周企业经营周报", "get_dashboard_metrics, analyze_project_risks, generate_weekly_report"],
            ["ticket-routing", "请把 4 号紧急工单自动派单", "route_ticket"],
        ],
        [3.2, 5.8, 6.5],
    )
    add_table(
        document,
        ["验证项", "Windows 结果", "WSL 结果", "说明"],
        [
            ["compileall", "通过", "通过", "检查 app、scripts、tests.py 语法"],
            ["tests.py", "通过", "通过", "覆盖鉴权、看板、Agent API 和密码哈希"],
            ["benchmark_agent.py", "3/3 通过", "3/3 通过", "工具调用命中率与关键词命中率达标"],
            ["smoke_web.py", "通过", "通过", "临时启动 Web 服务并自动关闭"],
            ["报告结构", "通过", "不适用", "DOCX 封面字段、TOC 域、标题导航均通过校验"],
        ],
        [3.0, 2.5, 2.5, 7.2],
    )
    add_screenshot_placeholder(document, "图 5-1 Demo 截图/录屏位置", "提交最终版前建议放入登录页、智能助手生成周报、工具调用轨迹、工单自动派单、审计日志页面截图，或在 docs/ 放置录屏链接说明。")

    document.add_heading("六、系统升级与扩展", level=1)
    for text in [
        "将 deterministic planner 替换为 LangGraph 状态图，加入条件分支、失败重试和人工确认节点。",
        "强化免费 DeepSeek 兼容 API、本地模型和 LangGraph 接入，用真实 Function Calling 生成工具调用计划。",
        "将 knowledge_base 表替换为向量数据库，实现语义检索和文档级 Agentic RAG。",
        "将 SQLite 替换为 PostgreSQL/MySQL，并接入企业统一身份认证。",
        "引入 OpenTelemetry、LangSmith 或自建 LLMOps 平台，实现 tracing、token 统计和质量评估。",
        "通过 Docker Compose 或云服务器部署提供可访问 URL，争取部署加分项。",
    ]:
        bullet(document, text)

    document.add_heading("七、课程总结", level=1)
    paragraph(document, "本项目最大的收获是：企业级应用的 Agent 改造不是简单加一个聊天框，而是要把业务系统中的数据、流程、权限和规则封装为可控工具，再让 Agent 在规格约束下编排工具完成业务目标。")
    paragraph(document, "通过本项目可以看到，SDD 方法使需求、架构、API、测试和报告形成闭环；Agentic AI 则让软件从“被动响应页面操作”升级为“主动完成业务目标”。开发者的角色也从单纯代码编写者转向规格制定者、工具设计者、状态管理者和行为评估者。")
    paragraph(document, "课程建议方面，希望后续能增加更多企业真实场景的 Agent 改造案例，并提供统一的 MCP、LangGraph 和 LLMOps 实验模板，帮助学生在有限时间内更深入比较不同 Agent 架构的优缺点。")

    document.add_heading("附录 A：评分标准对照", level=1)
    add_table(
        document,
        ["评分项", "分值", "本文档/项目对应内容"],
        [
            ["选题与设计思想", "20", "第一章说明原始系统痛点、项目价值和技术路线"],
            ["Specs 规格设计", "20", "第二章与 docs/specs/ 三份规格文档"],
            ["系统架构与设计", "15", "第三章架构图、Agent 交互流程和数据流设计"],
            ["关键实现与代码", "15", "第四章展示 Agent 核心循环、工具定义、配置与截图位置"],
            ["测试与评估", "10", "第五章覆盖功能测试、Benchmark、Web Smoke 与 Demo 截图位置"],
            ["升级扩展设想", "10", "第六章说明 LangGraph、向量库、LLMOps、云部署等演进路径"],
            ["课程总结", "10", "第七章总结工程思维转变和课程建议"],
        ],
        [4.0, 2.0, 9.5],
    )
    add_table(
        document,
        ["加分项", "分值", "项目体现"],
        [
            ["融合 MCP 协议 / Agentic RAG", "+3", "scripts/mcp_server.py 与 knowledge_base + search_knowledge"],
            ["云服务器部署并提供 URL", "+3", "已提供 Dockerfile/docker-compose.yml，后续可部署到云服务器"],
            ["课堂分享展示", "+2", "docs/demo_guide.md 提供 5 分钟展示脚本"],
            ["达到生产级水平", "+3", "错误回退、RBAC、密码哈希、审计日志、Agent tracing、CI 和 smoke test"],
        ],
        [4.2, 2.0, 9.3],
    )

    document.add_heading("附录 B：时间节点与提交说明", level=1)
    add_table(
        document,
        ["阶段", "时间", "交付物", "当前项目状态"],
        [
            ["Milestone 1: Proposal", "第 13 周（~06.01）", "设计文档、架构图、GitHub 仓库初始化、Spec 初稿", "docs/specs 已提供 Product/Architecture/API Spec"],
            ["Milestone 2: MVP", "第 14 周（~06.08）", "核心闭环 Demo 推送至 GitHub（tag: v0.1）", "Agent 周报、风险分析、派单、审批闭环可运行"],
            ["Final Demo Day", "第 16 周（~06.22）", "现场演示 5 分钟 + 答辩 3 分钟", "docs/demo_guide.md 已提供演示脚本"],
            ["最终提交截止", "2026 年 6 月 22 日 23:00", "GitHub 仓库最终版本（含完整文档和代码）", "仓库结构、README、LICENSE、报告、测试脚本已完成"],
        ],
        [3.4, 3.0, 5.0, 4.0],
    )
    paragraph(document, "最终提交方式以 GitHub cs599-project 仓库为唯一交付物；如果仓库为 Private，需要添加 qxr777 为 Collaborator；如果仓库为 Public，应保留 LICENSE 文件。打印版存档可使用 docs/CS599_大作业报告.pdf 或本 Word 文档导出的 PDF。")

    section = document.sections[-1]
    footer = section.footer.paragraphs[0]
    footer.text = ""
    add_page_number(footer)
    return document


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
